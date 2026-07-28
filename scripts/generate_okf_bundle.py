#!/usr/bin/env python3
"""Generate an OKF v0.2 knowledge bundle from any source folder.

Spec: https://github.com/GoogleCloudPlatform/knowledge-catalog/blob/main/okf/SPEC.md

Every concept is a thin markdown file whose `resource` field points at the
real source file via a file:// URI. Nothing is copied into the bundle.
Buckets are inferred from the source folder's own top-level subfolder names
(loose root-level files land in a "misc" bucket) rather than any fixed
domain taxonomy, so this works on an arbitrary folder of mixed documents.

Usage:
    python3 generate_okf_bundle.py --source <path> --dest <path> [--exclude "pat1,pat2"]
"""
import argparse
import csv
import fnmatch
import os
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path

import yaml

RUN_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()

EXCLUDED_DIRS = {".git", "__pycache__", "node_modules"}
EXCLUDED_FILE_EXACT = {".DS_Store"}
EXCLUDED_FILE_SUBSTR = ("secret", "credential", "password")
EXCLUDED_EXTS = {".mp4", ".mov", ".avi", ".env"}


def _matches_extra(relpath: Path, extra_patterns) -> bool:
    rel_str = str(relpath)
    return any(fnmatch.fnmatch(rel_str, pat) or fnmatch.fnmatch(relpath.name, pat) for pat in extra_patterns)


def is_dir_excluded(relpath: Path, extra_patterns) -> bool:
    # Directory-level pruning is deliberately narrow (exact names + explicit
    # --exclude patterns only). A content-style substring filter here (e.g.
    # "secret") would prune an entire directory just because it's *named*
    # "secrets/", silently swallowing everything inside without it ever being
    # counted as scanned or skipped.
    if relpath.name in EXCLUDED_DIRS:
        return True
    return _matches_extra(relpath, extra_patterns)


def is_file_excluded(relpath: Path, extra_patterns) -> bool:
    if relpath.name in EXCLUDED_FILE_EXACT:
        return True
    lower = relpath.name.lower()
    if any(s in lower for s in EXCLUDED_FILE_SUBSTR):
        return True
    if relpath.suffix.lower() in EXCLUDED_EXTS:
        return True
    return _matches_extra(relpath, extra_patterns)


def file_uri(path: Path) -> str:
    return path.resolve().as_uri()


def slugify(name: str) -> str:
    stem = Path(name).stem
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", stem).strip("-").lower()
    return slug or "untitled"


def extract_date_key(name: str):
    m = re.search(r"(\d{8})", name)
    if m:
        return m.group(1)
    m = re.search(r"(\d{4}[-_]\d{2}[-_]\d{2})", name)
    if m:
        return re.sub(r"[-_]", "", m.group(1))
    return None


# ---------------------------------------------------------------------------
# Per-type extractors: each returns (description, body_markdown)
# ---------------------------------------------------------------------------

def extract_docx(path: Path):
    try:
        import docx
        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"Could not extract text ({e}).", ""
    excerpt = text.strip()[:1200]
    description = " ".join(excerpt.split())[:180]
    body = f"## Extracted Excerpt\n\n{excerpt}\n\n_(truncated — see source document)_"
    return description, body


def extract_pdf(path: Path):
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        pages = reader.pages[:2]
        text = "\n".join(p.extract_text() or "" for p in pages)
    except Exception as e:
        return f"Could not extract text ({e}).", ""
    excerpt = text.strip()[:1200]
    description = " ".join(excerpt.split())[:180]
    body = (
        f"## Extracted Excerpt (first pages)\n\n{excerpt}\n\n"
        f"_(truncated — {page_count} total pages in source)_"
    )
    return description, body


def extract_xlsx(path: Path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    except Exception as e:
        return f"Could not read workbook ({e}).", ""
    lines = ["## Schema", "", "| Sheet | Columns |", "|---|---|"]
    for sheet_name in wb.sheetnames[:8]:
        ws = wb[sheet_name]
        try:
            header = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        except StopIteration:
            header = ()
        cols = ", ".join(str(c) for c in header if c is not None)[:300]
        lines.append(f"| {sheet_name} | {cols or '(empty)'} |")
    description = f"Spreadsheet with {len(wb.sheetnames)} sheet(s): {', '.join(wb.sheetnames[:5])}."
    wb.close()
    return description, "\n".join(lines)


def extract_duckdb_schema(path: Path):
    duckdb_bin = "duckdb"
    # SHOW TABLES only lists the default `main` schema. Query
    # information_schema.tables across all schemas so warehouses using named
    # schemas (e.g. xero_data, audit) are fully introspected.
    tables_out = subprocess.run(
        [duckdb_bin, str(path), "-csv", "-c",
         "SELECT table_schema, table_name FROM information_schema.tables "
         "WHERE table_schema NOT IN ('information_schema', 'pg_catalog') "
         "ORDER BY table_schema, table_name;"],
        capture_output=True, text=True, timeout=60,
    )
    rows = [r for r in tables_out.stdout.splitlines()[1:] if r.strip()]
    tables = []  # list of (schema, table)
    for row in rows:
        cols = next(csv.reader([row]))
        if len(cols) >= 2:
            tables.append((cols[0], cols[1]))

    lines = ["## Schema", ""]
    for schema, table in tables:
        qualified = f"{schema}.{table}"
        lines.append(f"### Table: {qualified}")
        lines.append("")
        lines.append("| Column | Type |")
        lines.append("|---|---|")
        desc = subprocess.run(
            [duckdb_bin, str(path), "-csv", "-c", f'DESCRIBE "{schema}"."{table}";'],
            capture_output=True, text=True, timeout=60,
        )
        desc_rows = [r for r in desc.stdout.splitlines()[1:] if r.strip()]
        for row in desc_rows:
            cols = next(csv.reader([row]))
            if len(cols) >= 2:
                lines.append(f"| {cols[0]} | {cols[1]} |")
        lines.append("")
    table_names = [f"{s}.{t}" for s, t in tables]
    description = f"DuckDB warehouse with {len(tables)} table(s): {', '.join(table_names[:6])}."
    return description, "\n".join(lines)


def extract_csv_folder(folder: Path):
    lines = ["## Schema (per file)", "", "| File | Columns | Rows |", "|---|---|---|"]
    csv_files = sorted(folder.glob("*.csv"))
    for csv_path in csv_files:
        try:
            with open(csv_path, newline="", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                header = next(reader, [])
                row_count = sum(1 for _ in reader)
        except Exception:
            header, row_count = [], "?"
        cols = ", ".join(header)[:250]
        lines.append(f"| {csv_path.name} | {cols} | {row_count} |")
    description = f"{len(csv_files)} CSV file(s) in {folder.name}/."
    return description, "\n".join(lines)


def extract_text_excerpt(path: Path, max_chars=1200):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(max_chars * 4)
    except Exception as e:
        return f"Could not read file ({e}).", ""
    excerpt = text.strip()[:max_chars]
    description = " ".join(excerpt.split())[:180] or f"Reference file: {path.name}."
    body = f"## Excerpt\n\n```\n{excerpt}\n```\n\n_(truncated)_"
    return description, body


def extract_image_stub(path: Path):
    return f"Image asset: {path.name}.", "No OCR performed; filename is the only available context."


# ---------------------------------------------------------------------------
# Concept model + writer
# ---------------------------------------------------------------------------

class Concept:
    def __init__(self, bundle_relpath, type_, title, description, resource_path, tags, body, related=None):
        self.bundle_relpath = bundle_relpath
        self.type = type_
        self.title = title
        self.description = description
        self.resource_path = resource_path  # Path or None
        self.tags = tags
        self.body = body
        self.related = related or []  # list of (title, bundle_relpath)


def write_concept(bundle_root: Path, generated_by: str, concept: Concept):
    dest = bundle_root / concept.bundle_relpath
    dest.parent.mkdir(parents=True, exist_ok=True)
    frontmatter = {
        "type": concept.type,
        "title": concept.title,
        "description": concept.description,
    }
    if concept.resource_path is not None:
        frontmatter["resource"] = file_uri(concept.resource_path)
    if concept.tags:
        frontmatter["tags"] = concept.tags
    frontmatter["generated"] = {"by": generated_by, "at": RUN_AT}

    yaml_block = yaml.dump(frontmatter, sort_keys=False, allow_unicode=True, default_flow_style=False)
    body = concept.body
    if concept.related:
        related_lines = "\n".join(f"- [{t}](/{r})" for t, r in concept.related)
        body += f"\n\n## Related\n\n{related_lines}"
    content = f"---\n{yaml_block}---\n\n# {concept.title}\n\n{body}\n"
    dest.write_text(content, encoding="utf-8")


# ---------------------------------------------------------------------------
# Classification — generic: bucket = source's own top-level subfolder name
# ---------------------------------------------------------------------------

EXT_TYPE = {
    ".xlsx": ("Spreadsheet", "spreadsheet"),
    ".xls": ("Spreadsheet", "spreadsheet"),
    ".docx": ("Document", "document"),
    ".doc": ("Document", "document"),
    ".pdf": ("Document", "document"),
    ".png": ("Media Asset", "image"),
    ".jpg": ("Media Asset", "image"),
    ".jpeg": ("Media Asset", "image"),
    ".webp": ("Media Asset", "image"),
    ".svg": ("Media Asset", "image"),
    ".md": ("Note", "note"),
    ".yaml": ("Reference", "config"),
    ".yml": ("Reference", "config"),
    ".json": ("Reference", "config"),
    ".txt": ("Reference", "text"),
    ".html": ("Reference", "text"),
    ".log": ("Reference", "log"),
}


def classify(relpath: Path):
    """Return (bucket, type_str, tags) or None to skip (handled in an aggregate pass)."""
    ext = relpath.suffix.lower()
    if ext in (".duckdb", ".csv"):
        return None  # handled by aggregate passes

    top = relpath.parts[0] if len(relpath.parts) > 1 else None
    bucket = re.sub(r"[^a-zA-Z0-9]+", "-", top).strip("-").lower() if top else "misc"
    bucket = bucket or "misc"

    type_, tag = EXT_TYPE.get(ext, ("Reference", "misc"))
    tags = [bucket, tag]
    return (bucket, type_, tags)


def build_generic_concept(src_path: Path, relpath: Path, bucket: str, type_: str, tags: list):
    ext = relpath.suffix.lower()
    if ext in (".docx", ".doc"):
        description, body = extract_docx(src_path)
    elif ext == ".pdf":
        description, body = extract_pdf(src_path)
    elif ext in (".xlsx", ".xls"):
        description, body = extract_xlsx(src_path)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".svg"):
        description, body = extract_image_stub(src_path)
    else:
        description, body = extract_text_excerpt(src_path)

    date_key = extract_date_key(relpath.name)
    # Slug from the full relative path (parent dirs + stem + extension), not just
    # the bare filename, so same-name files in different subfolders or same-stem
    # twins with different extensions never collide and silently overwrite each
    # other. The extension token also guarantees a source named "index.html"
    # never produces the reserved "index.md" filename.
    base_slug = slugify(relpath.name)
    ext_token = ext.lstrip(".") or "file"
    parent_prefix = "-".join(
        re.sub(r"[^a-zA-Z0-9]+", "-", part).strip("-").lower() for part in relpath.parent.parts
    )
    full_slug = "-".join(p for p in [parent_prefix, base_slug, ext_token] if p)
    if date_key:
        fname = f"{date_key}-{full_slug}.md"
    else:
        fname = f"{full_slug}.md"
    if fname in ("index.md", "log.md"):
        fname = f"{full_slug}-doc.md"
    bundle_relpath = f"{bucket}/{fname}"

    title = Path(relpath.name).stem.replace("_", " ").replace("-", " ").strip()
    return Concept(bundle_relpath, type_, title, description, src_path, tags, body), date_key


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", required=True, help="Source folder to bundle")
    parser.add_argument("--dest", required=True, help="Destination bundle folder (created/overwritten)")
    parser.add_argument(
        "--exclude", default="",
        help="Comma-separated glob patterns (matched against filename or relative path) to skip",
    )
    args = parser.parse_args()

    source_root = Path(args.source).expanduser().resolve()
    bundle_root = Path(args.dest).expanduser().resolve()
    if not source_root.is_dir():
        raise SystemExit(f"Source folder does not exist: {source_root}")
    extra_patterns = [p.strip() for p in args.exclude.split(",") if p.strip()]

    generated_by = f"process:okf-ingest:{bundle_root.name}"

    if bundle_root.exists():
        print(f"Removing existing bundle at {bundle_root} before regenerating...")
        shutil.rmtree(bundle_root)
    bundle_root.mkdir(parents=True)

    all_concepts = []
    date_index = {}  # date_key -> list of (title, bundle_relpath)
    source_file_count = 0
    skipped_count = 0
    duckdb_paths = []
    csv_dirs = {}  # parent relpath -> list of csv file relpaths

    for root, dirs, files in os.walk(source_root):
        rel_root = Path(root).relative_to(source_root)
        pruned = [d for d in dirs if is_dir_excluded((rel_root / d) if str(rel_root) != "." else Path(d), extra_patterns)]
        for d in pruned:
            drelpath = (rel_root / d) if str(rel_root) != "." else Path(d)
            skipped_count += sum(1 for _ in (source_root / drelpath).rglob("*") if _.is_file())
        dirs[:] = [d for d in dirs if d not in pruned]
        for fname in files:
            relpath = (rel_root / fname) if str(rel_root) != "." else Path(fname)
            src_path = source_root / relpath
            if is_file_excluded(relpath, extra_patterns):
                skipped_count += 1
                continue
            source_file_count += 1

            ext = relpath.suffix.lower()
            if ext == ".duckdb":
                duckdb_paths.append(relpath)
                continue
            if ext == ".csv":
                csv_dirs.setdefault(str(relpath.parent), []).append(relpath)
                continue

            classification = classify(relpath)
            if classification is None:
                continue
            bucket, type_, tags = classification
            concept, date_key = build_generic_concept(src_path, relpath, bucket, type_, tags)
            all_concepts.append(concept)
            if date_key:
                date_index.setdefault(date_key, []).append((concept.title, concept.bundle_relpath))

    # Cross-link concepts that share the same detected date (e.g. same-day meeting artifacts)
    for concept in all_concepts:
        date_key = extract_date_key(Path(concept.bundle_relpath).name)
        if not date_key:
            continue
        siblings = [(t, p) for t, p in date_index.get(date_key, []) if p != concept.bundle_relpath]
        concept.related = siblings

    # --- Aggregate: one concept per directory containing CSV files ---
    for dirname, csv_relpaths in sorted(csv_dirs.items()):
        folder = source_root / dirname
        first_relpath = csv_relpaths[0]
        top = first_relpath.parts[0] if len(first_relpath.parts) > 1 else None
        bucket = re.sub(r"[^a-zA-Z0-9]+", "-", top).strip("-").lower() if top else "misc"
        description, body = extract_csv_folder(folder)
        slug = slugify(dirname if dirname != "." else folder.name)
        concept = Concept(
            f"{bucket}/{slug}-csv.md", "Data Export", folder.name, description,
            folder, [bucket, "csv"], body,
        )
        all_concepts.append(concept)

    # --- Aggregate: one concept per .duckdb file found ---
    for relpath in duckdb_paths:
        src_path = source_root / relpath
        top = relpath.parts[0] if len(relpath.parts) > 1 else None
        bucket = re.sub(r"[^a-zA-Z0-9]+", "-", top).strip("-").lower() if top else "database"
        try:
            description, body = extract_duckdb_schema(src_path)
        except Exception as e:
            description, body = f"Could not introspect schema ({e}).", ""
        slug = slugify(relpath.name)
        concept = Concept(
            f"{bucket}/{slug}-duckdb.md", "Database", relpath.stem, description,
            src_path, [bucket, "duckdb"], body,
        )
        all_concepts.append(concept)

    # Write all concepts
    for concept in all_concepts:
        write_concept(bundle_root, generated_by, concept)

    # --- index.md per bucket directory + root ---
    by_dir = {}
    for concept in all_concepts:
        dirname = str(Path(concept.bundle_relpath).parent)
        by_dir.setdefault(dirname, []).append(concept)

    for dirname, concepts in by_dir.items():
        lines = [f"# {dirname.replace('-', ' ').title()}", ""]
        for c in sorted(concepts, key=lambda c: c.bundle_relpath):
            fname = Path(c.bundle_relpath).name
            lines.append(f"* [{c.title}]({fname}) - {c.description}")
        (bundle_root / dirname / "index.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    root_lines = [
        "---",
        'okf_version: "0.2"',
        "---",
        "",
        f"# {bundle_root.name} Knowledge Bundle",
        "",
        f"OKF v0.2 bundle generated from `{source_root}`. All `resource:` fields are",
        "`file://` URIs pointing at the original files on this machine (point-in-place,",
        "nothing copied). These paths will not resolve on another machine unless the",
        "source folder is synced there too.",
        "",
    ]
    for dirname in sorted(by_dir):
        root_lines.append(f"# {dirname.replace('-', ' ').title()}")
        root_lines.append("")
        root_lines.append(f"* [{dirname}/]({dirname}/) - {len(by_dir[dirname])} concept(s)")
        root_lines.append("")
    (bundle_root / "index.md").write_text("\n".join(root_lines), encoding="utf-8")

    log_lines = [
        "# Directory Update Log",
        "",
        f"## {RUN_AT[:10]}",
        f"* **Initialization**: Generated {len(all_concepts)} concepts from `{source_root}` "
        f"({source_file_count} source files scanned, {skipped_count} excluded).",
    ]
    (bundle_root / "log.md").write_text("\n".join(log_lines) + "\n", encoding="utf-8")

    print(f"Scanned {source_file_count} files, skipped {skipped_count} excluded files.")
    print(f"Wrote {len(all_concepts)} concepts into {bundle_root}")


if __name__ == "__main__":
    main()
